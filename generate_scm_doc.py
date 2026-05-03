#!/usr/bin/env python3
"""Generate the SCM tables as a clean Word document."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path("SCM_Tables.docx")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)

    return table


def main():
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    add_heading(doc, "Synthetic Data Generation: Structural Causal Model", level=1)

    # ── Explanatory paragraph ──────────────────────────────────────────────────
    add_paragraph(doc,
        "To generate synthetic therapy narratives with a known ground truth, we specify a "
        "structural causal model (SCM) consistent with the ground truth DAG. Each variable "
        "is binary. Endogenous variables are generated via a logistic SCM: a latent "
        "continuous score X* is computed as a linear combination of parent values plus an "
        "independent logistic noise term εᵢ ~ Logistic(0, 1), and the observed "
        "binary variable is X = 1(X* > 0). The root variable early_adversity has no parents "
        "and is drawn directly from a Bernoulli distribution. In addition to the six DAG "
        "variables, each document is assigned a set of noise variables — life "
        "circumstances that appear in the narrative but have no causal connection to "
        "depression. Noise variables are sampled independently of the SCM: the number of "
        "noise variables per document is drawn uniformly from {0, 1, 2}, and the selected "
        "variables are drawn without replacement from a fixed pool, with each variable in "
        "the pool having equal probability of selection."
    )

    doc.add_paragraph()

    # ── Table 1: Structural Equations ─────────────────────────────────────────
    add_heading(doc, "Table 1 — Structural Equations", level=2)
    add_paragraph(doc,
        "Each variable follows: X* = β₀ + terms + εᵢ, "
        "X = 1(X* > 0), where εᵢ ~ Logistic(0, 1) independently."
    )

    eq_headers = ["Variable", "β₀ (intercept)", "β₁", "β₂", "β₁₂ (interaction)"]
    eq_rows = [
        ["early_adversity",       "−0.847", "—",                                    "—",                                        "—"],
        ["chronic_stress",        "−1.099", "+1.718 · early_adversity",             "—",                                        "—"],
        ["emotion_dysregulation", "−2.197", "+1.578 · early_adversity",             "+1.350 · chronic_stress",                  "−0.112 · early_adversity · chronic_stress"],
        ["social_withdrawal",     "−1.735", "+1.935 · chronic_stress",              "—",                                        "—"],
        ["rumination",            "−2.197", "+1.792 · emotion_dysregulation",       "+1.578 · social_withdrawal",               "−0.229 · emotion_dysregulation · social_withdrawal"],
        ["depression",            "−2.944", "+2.744 · rumination",                 "+1.209 · early_adversity",                 "−0.162 · rumination · early_adversity"],
    ]
    make_table(doc, eq_headers, eq_rows)

    doc.add_paragraph()

    # ── Table 2: Conditional Probability Table ─────────────────────────────────
    add_heading(doc, "Table 2 — Conditional Probability Table", level=2)

    cpt_headers = ["Variable", "early_adversity", "chronic_stress", "emotion_dysregulation", "social_withdrawal", "rumination", "P(= 1)"]
    cpt_rows = [
        ["early_adversity",       "—", "—", "—", "—", "—", "0.30"],
        ["chronic_stress",        "0",      "—", "—", "—", "—", "0.25"],
        ["chronic_stress",        "1",      "—", "—", "—", "—", "0.65"],
        ["emotion_dysregulation", "0", "0", "—", "—", "—", "0.10"],
        ["emotion_dysregulation", "1", "0", "—", "—", "—", "0.35"],
        ["emotion_dysregulation", "0", "1", "—", "—", "—", "0.30"],
        ["emotion_dysregulation", "1", "1", "—", "—", "—", "0.65"],
        ["social_withdrawal",     "—", "0", "—", "—", "—", "0.15"],
        ["social_withdrawal",     "—", "1", "—", "—", "—", "0.55"],
        ["rumination",            "—", "—", "0", "0", "—", "0.10"],
        ["rumination",            "—", "—", "1", "0", "—", "0.40"],
        ["rumination",            "—", "—", "0", "1", "—", "0.35"],
        ["rumination",            "—", "—", "1", "1", "—", "0.72"],
        ["depression",            "0", "—", "—", "—", "0", "0.05"],
        ["depression",            "0", "—", "—", "—", "1", "0.45"],
        ["depression",            "1", "—", "—", "—", "0", "0.15"],
        ["depression",            "1", "—", "—", "—", "1", "0.70"],
    ]
    make_table(doc, cpt_headers, cpt_rows)

    doc.add_paragraph()

    # ── Table 3: Noise Variable Pool ───────────────────────────────────────────
    add_heading(doc, "Table 3 — Noise Variable Pool", level=2)
    add_paragraph(doc,
        "Noise variables are sampled independently of the DAG. The count per document "
        "is drawn from Uniform{0, 1, 2}. Each variable in the pool has equal selection "
        "probability."
    )

    noise_headers = ["Noise Variable", "Description"]
    noise_rows = [
        ["relationship_conflict", "Interpersonal tension with a partner, friend, or family member"],
        ["diet_change",           "Recent changes to eating habits or nutrition"],
        ["housing_change",        "Moving home or changes to living situation"],
        ["work_promotion",        "Positive career development at work"],
        ["new_hobby",             "Taking up a new activity or interest"],
        ["financial_planning",    "Actively managing savings or budget"],
        ["travel",                "Recent or upcoming trips"],
        ["physical_health_issue", "Minor physical ailment unrelated to mental health"],
        ["family_event",          "A significant event involving family (wedding, birth, etc.)"],
        ["career_change",         "Switching jobs or professional direction"],
    ]
    make_table(doc, noise_headers, noise_rows)

    doc.save(OUT)
    print(f"Saved → {OUT}")


if __name__ == "__main__":
    main()
